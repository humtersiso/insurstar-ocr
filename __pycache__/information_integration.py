#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word模板整合測試
測試從OCR到Word模板生成的完整流程
包含Word模板修復功能
"""

import os
import json
import glob
from datetime import datetime
from docx import Document
import re
from word_template_processor import WordTemplateProcessor

# 浮水印與特殊欄位對應
WATERMARK_MAP = {
    '{{watermark_name_blue}}': 'assets/watermark/watermark_name_blue.png',
    '{{watermark_company_blue}}': 'assets/watermark/watermark_company_blue.png',
    '{{PCN}}': 'BB2H699299',
}

def replace_special_tags(text):
    for k, v in WATERMARK_MAP.items():
        text = text.replace(k, v)
    return text

class WordTemplateFixer:
    """Word模板修復器"""
    
    @staticmethod
    def fix_template_syntax(original_path: str, fixed_path: str) -> bool:
        """修復模板中的語法錯誤"""
        print(f"🔧 修復Word模板語法錯誤")
        print(f"   原始: {original_path}")
        print(f"   修復: {fixed_path}")
        
        if not os.path.exists(original_path):
            print(f"❌ 原始模板檔案不存在: {original_path}")
            return False
        
        try:
            # 載入原始模板
            doc = Document(original_path)
            
            # 修復表格中的語法錯誤
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # 修復空的模板標記
                            if "{{}}" in paragraph.text:
                                print(f"   🔧 修復空標記: {paragraph.text[:50]}...")
                                paragraph.text = paragraph.text.replace("{{}}", "{{gender}}")
                            
                            # 修復其他可能的語法問題
                            text = paragraph.text
                            open_count = text.count("{{")
                            close_count = text.count("}}")
                            
                            if open_count != close_count:
                                print(f"   ⚠️ 標記不匹配: {text[:50]}...")
                                # 移除有問題的標記
                                text = re.sub(r'\{\{[^}]*$', '', text)
                                text = re.sub(r'^\}[^}]*\}\}', '', text)
                                paragraph.text = text
            
            # 修復段落中的語法錯誤
            for paragraph in doc.paragraphs:
                if "{{}}" in paragraph.text:
                    print(f"   🔧 修復段落空標記: {paragraph.text[:50]}...")
                    paragraph.text = paragraph.text.replace("{{}}", "{{gender}}")
            
            # 儲存修復後的模板
            doc.save(fixed_path)
            print(f"✅ 修復完成: {fixed_path}")
            
            return True
            
        except Exception as e:
            print(f"❌ 修復失敗: {str(e)}")
            return False

def find_latest_ocr_result():
    """尋找最新的OCR結果檔案"""
    print("🔍 尋找最新的OCR結果...")
    
    # 優先搜尋原始的OCR結果檔案（UUID命名）
    uuid_pattern = "ocr_results/*_gemini_ocr_output.json"
    uuid_files = glob.glob(uuid_pattern)
    
    if uuid_files:
        # 按修改時間排序，取最新的
        latest_file = max(uuid_files, key=os.path.getmtime)
        print(f"✅ 找到最新原始OCR結果: {latest_file}")
        return latest_file
    
    # 如果沒有原始OCR結果，搜尋處理後的資料檔案
    json_pattern = "ocr_results/*_data.json"
    json_files = glob.glob(json_pattern)
    
    if json_files:
        # 按修改時間排序，取最新的
        latest_file = max(json_files, key=os.path.getmtime)
        print(f"✅ 找到最新處理後OCR結果: {latest_file}")
        return latest_file
    
    print("❌ 未找到OCR結果檔案")
    return None

def load_ocr_data(json_path: str):
    """載入OCR資料"""
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f"📊 載入OCR資料成功:")
        print(f"   - 檔案: {json_path}")
        print(f"   - 欄位數量: {len(data)}")
        
        # 顯示關鍵欄位
        key_fields = ['insurance_company', 'insured_person', 'policyholder', 'vehicle_type', 'gender']
        print(f"   - 關鍵欄位:")
        for field in key_fields:
            value = data.get(field, '無資料')
            print(f"     * {field}: {value}")
        
        return data
        
    except Exception as e:
        print(f"❌ 載入OCR資料失敗: {str(e)}")
        return None

def select_ocr_json_file():
    """讓使用者選擇ocr_results資料夾下的OCR JSON檔案"""
    import glob
    import os
    json_files = glob.glob('ocr_results/*.json')
    if not json_files:
        print('❌ ocr_results 資料夾下沒有 json 檔案')
        return None
    print('\n請選擇要測試的 OCR JSON 檔案：')
    for idx, f in enumerate(json_files):
        print(f"  [{idx+1}] {os.path.basename(f)}")
    while True:
        try:
            sel = int(input('請輸入檔案編號：'))
            if 1 <= sel <= len(json_files):
                return json_files[sel-1]
        except Exception:
            pass
        print('輸入錯誤，請重新輸入。')

def test_complete_workflow():
    """測試完整的工作流程"""
    print("🔄 Word模板整合測試")
    print("=" * 60)
    
    try:
        # 步驟1: 修復Word模板
        print("\n🔧 步驟1: 修復Word模板...")
        original_template = "assets/templates/財產分析書.docx"
        fixed_template = "assets/templates/財產分析書_fixed.docx"
        
        if not WordTemplateFixer.fix_template_syntax(original_template, fixed_template):
            print("❌ Word模板修復失敗")
            return False
        
        # 步驟2: 選擇並載入OCR結果
        print("\n📄 步驟2: 載入OCR結果...")
        ocr_json_path = select_ocr_json_file()
        if not ocr_json_path:
            print("❌ 未選擇OCR結果檔案")
            return False
        ocr_data = load_ocr_data(ocr_json_path)
        if not ocr_data:
            print("❌ 載入OCR資料失敗")
            return False
        
        # 步驟3: 使用修復後的模板生成Word
        print("\n📄 步驟3: 生成Word檔案...")
        word_processor = WordTemplateProcessor(fixed_template)
        
        # 生成Word檔案
        output_path = word_processor.fill_template(ocr_data)
        
        if output_path:
            print(f"✅ Word檔案生成成功: {output_path}")
            
            # 檢查檔案是否存在
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"   - 檔案大小: {file_size:,} bytes")
                # 新增：檢查保期勾選欄位
                processed = word_processor.process_ocr_data(ocr_data)
                compulsory_check = processed.get("check_compulsory_insurance_period", "")
                optional_check = processed.get("check_optional_insurance_period", "")
                print(f"   - 強制險保期勾選: {compulsory_check}")
                print(f"   - 任意險保期勾選: {optional_check}")
                if ocr_data.get("compulsory_insurance_period"):
                    assert compulsory_check == "☑ ", "強制險保期有資料時應打勾"
                else:
                    assert compulsory_check == "□", "強制險保期無資料時應空白"
                if ocr_data.get("optional_insurance_period"):
                    assert optional_check == "☑ ", "任意險保期有資料時應打勾"
                else:
                    assert optional_check == "□", "任意險保期無資料時應空白"
            else:
                print("❌ 生成的檔案不存在")
                return False
        else:
            print("❌ Word檔案生成失敗")
            return False
        
        # 步驟4: 儲存完整處理資料
        print("\n💾 步驟4: 儲存完整處理資料...")
        json_output_path = output_path.replace('.docx', '_complete_data.json')
        word_processor.save_processed_data(ocr_data, json_output_path)
        
        print(f"✅ 完整資料已儲存: {json_output_path}")
        
        # 顯示摘要
        print("\n📋 處理摘要:")
        print(f"   - 原始模板: {original_template}")
        print(f"   - 修復模板: {fixed_template}")
        print(f"   - OCR資料: {ocr_json_path}")
        print(f"   - Word檔案: {output_path}")
        print(f"   - 完整資料: {json_output_path}")
        
        print("\n🎉 整合測試完成！")
        return True
        
    except Exception as e:
        print(f"❌ 測試失敗: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def test_api_integration():
    """測試API整合"""
    print("\n🌐 API整合測試")
    print("=" * 60)
    
    try:
        import requests
        
        # 測試API端點
        base_url = "http://localhost:5000"
        
        # 健康檢查
        print("🔍 健康檢查...")
        response = requests.get(f"{base_url}/api/health")
        if response.status_code == 200:
            print("✅ API服務正常")
        else:
            print(f"❌ API服務異常: {response.status_code}")
            return False
        
        # 載入測試資料
        ocr_json_path = find_latest_ocr_result()
        if not ocr_json_path:
            print("❌ 未找到OCR結果檔案")
            return False
        
        with open(ocr_json_path, 'r', encoding='utf-8') as f:
            test_ocr_data = json.load(f)
        
        # 測試Word模板API
        print("\n📄 測試Word模板API...")
        
        api_data = {
            "ocr_data": test_ocr_data
        }
        
        response = requests.post(
            f"{base_url}/api/generate-word-template",
            json=api_data,
            headers={'Content-Type': 'application/json'}
        )
        
        if response.status_code == 200:
            result = response.json()
            print("✅ Word模板API測試成功")
            print(f"   - 檔案: {result.get('word_filename')}")
            print(f"   - 下載連結: {result.get('download_url')}")
        else:
            print(f"❌ Word模板API測試失敗: {response.status_code}")
            print(f"   回應: {response.text}")
            return False
        
        print("\n🎉 API整合測試完成！")
        return True
        
    except Exception as e:
        print(f"❌ API測試失敗: {str(e)}")
        return False

def main():
    """主函數"""
    print("🧪 Word模板整合測試")
    print("=" * 80)
    
    # 測試完整工作流程
    workflow_success = test_complete_workflow()
    
    # 測試API整合（需要Flask服務運行）
    api_success = False
    try:
        api_success = test_api_integration()
    except:
        print("⚠️ API測試跳過（Flask服務未運行）")
    
    # 總結
    print("\n📊 測試結果總結")
    print("=" * 80)
    print(f"完整工作流程: {'✅ 成功' if workflow_success else '❌ 失敗'}")
    print(f"API整合測試: {'✅ 成功' if api_success else '❌ 失敗'}")
    
    if workflow_success:
        print("\n🎉 Word模板處理系統已準備就緒！")
        print("📝 使用方式:")
        print("   1. 直接使用 WordTemplateProcessor 類別")
        print("   2. 透過 Flask API: /api/generate-word-template")
        print("   3. 傳入OCR辨識結果，自動生成財產分析書")
        print("\n🔄 流程說明:")
        print("   1. 自動修復Word模板語法錯誤")
        print("   2. 從ocr_results資料夾讀取最新OCR結果")
        print("   3. 根據標記填入辨識資料")
        print("   4. 產生完整的Word檔案")
    else:
        print("\n⚠️ 系統需要進一步檢查")

if __name__ == "__main__":
    main() 