#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
詳細分析 Word 模板中的所有標記
檢查所有可能的位置：段落、表格、頁首頁尾、註解、文字方塊等
"""

import os
import re
from docx import Document
from docxtpl import DocxTemplate

def analyze_template_detailed(template_path):
    """詳細分析模板中的所有標記"""
    print(f"🔍 詳細分析模板: {template_path}")
    print("=" * 60)
    
    if not os.path.exists(template_path):
        print(f"❌ 檔案不存在: {template_path}")
        return
    
    try:
        # 使用 python-docx 載入文檔
        doc = Document(template_path)
        
        all_text = ""
        sections_found = []
        
        # 1. 分析段落
        print("\n📝 1. 段落分析:")
        print("-" * 30)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                all_text += para.text + "\n"
                # 檢查是否包含標記
                if "{{" in para.text or "}}" in para.text:
                    print(f"  段落 {i+1}: {para.text.strip()}")
        
        # 2. 分析表格
        print(f"\n📊 2. 表格分析 (共 {len(doc.tables)} 個表格):")
        print("-" * 30)
        for table_idx, table in enumerate(doc.tables):
            print(f"  表格 {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            all_text += para.text + "\n"
                            # 檢查是否包含標記
                            if "{{" in para.text or "}}" in para.text:
                                print(f"    表格{table_idx+1} 行{row_idx+1} 格{cell_idx+1} 段落{para_idx+1}: {para.text.strip()}")
        
        # 3. 分析頁首頁尾
        print(f"\n📄 3. 頁首頁尾分析:")
        print("-" * 30)
        for i, section in enumerate(doc.sections):
            print(f"  區段 {i+1}:")
            
            # 頁首
            if section.header:
                header_text = ""
                for para in section.header.paragraphs:
                    header_text += para.text + "\n"
                if header_text.strip():
                    all_text += header_text + "\n"
                    if "{{" in header_text or "}}" in header_text:
                        print(f"    頁首: {header_text.strip()}")
            
            # 頁尾
            if section.footer:
                footer_text = ""
                for para in section.footer.paragraphs:
                    footer_text += para.text + "\n"
                if footer_text.strip():
                    all_text += footer_text + "\n"
                    if "{{" in footer_text or "}}" in footer_text:
                        print(f"    頁尾: {footer_text.strip()}")
        
        # 4. 分析註解
        print(f"\n💬 4. 註解分析:")
        print("-" * 30)
        # 註解通常在 docx 的 XML 中，這裡簡化處理
        
        # 5. 分析文字方塊和圖形
        print(f"\n📦 5. 文字方塊/圖形分析:")
        print("-" * 30)
        # 文字方塊通常在 docx 的 XML 中，這裡簡化處理
        
        # 6. 使用 docxtpl 載入並分析
        print(f"\n🔧 6. DocxTemplate 分析:")
        print("-" * 30)
        try:
            docxtpl_doc = DocxTemplate(template_path)
            
            # 取得所有文字
            docxtpl_text = ""
            if docxtpl_doc.docx and docxtpl_doc.docx.paragraphs:
                for para in docxtpl_doc.docx.paragraphs:
                    docxtpl_text += para.text + "\n"
            
            if docxtpl_doc.docx and docxtpl_doc.docx.tables:
                for table in docxtpl_doc.docx.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                docxtpl_text += para.text + "\n"
            
            # 找出所有標記
            markers = re.findall(r'\{\{[^}]+\}\}', docxtpl_text)
            print(f"  DocxTemplate 找到 {len(markers)} 個標記:")
            for marker in sorted(set(markers)):
                print(f"    - {marker}")
            
            all_text += docxtpl_text
            
        except Exception as e:
            print(f"  DocxTemplate 分析失敗: {str(e)}")
        
        # 7. 綜合分析
        print(f"\n📋 7. 綜合分析結果:")
        print("-" * 30)
        
        # 找出所有可能的標記
        all_markers = re.findall(r'\{\{[^}]+\}\}', all_text)
        print(f"  總共找到 {len(all_markers)} 個標記:")
        for marker in sorted(set(all_markers)):
            print(f"    - {marker}")
        
        # 檢查特定標記
        target_markers = ['{{watermark_name_blue}}', '{{watermark_company_blue}}', '{{PCN}}']
        print(f"\n🎯 目標標記檢查:")
        for marker in target_markers:
            if marker in all_markers:
                print(f"  ✅ {marker} - 存在")
            else:
                print(f"  ❌ {marker} - 不存在")
        
        # 檢查被拆開的標記
        print(f"\n🔍 被拆開的標記檢查:")
        if "{{P" in all_text and "CN}}" in all_text:
            print(f"  ⚠️ 發現被拆開的 PCN 標記: {{P + CN}}")
        if "{{watermark_" in all_text and "name_blue}}" in all_text:
            print(f"  ⚠️ 發現被拆開的 watermark_name_blue 標記")
        if "{{watermark_" in all_text and "company_blue}}" in all_text:
            print(f"  ⚠️ 發現被拆開的 watermark_company_blue 標記")
        
        # 顯示包含標記的行
        print(f"\n📝 包含標記的完整上下文:")
        lines = all_text.split('\n')
        for i, line in enumerate(lines):
            if any(marker.replace('{{', '').replace('}}', '') in line for marker in target_markers):
                print(f"  第{i+1}行: {line.strip()}")
            elif "{{P" in line or "CN}}" in line:
                print(f"  第{i+1}行: {line.strip()}")
            elif "{{watermark_" in line:
                print(f"  第{i+1}行: {line.strip()}")
        
        return all_markers
        
    except Exception as e:
        print(f"❌ 分析失敗: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

def main():
    """主函數"""
    print("🔍 Word 模板詳細分析工具")
    print("=" * 80)
    
    # 分析三個模板
    templates = [
        "assets/templates/財產分析書.docx",
        "assets/templates/財產分析書_fixed.docx",
        "assets/templates/財產分析書_simple.docx"
    ]
    
    for template_path in templates:
        analyze_template_detailed(template_path)
        print("\n" + "="*80)

if __name__ == "__main__":
    main() 