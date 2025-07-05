#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 模板分析工具
詳細分析 Word 模板的內容結構，找出需要填入的位置
"""

from docx import Document
import os
from typing import Dict, List, Any

def analyze_word_template(template_path: str):
    """
    分析 Word 模板內容
    
    Args:
        template_path: Word 模板檔案路徑
    """
    print(f"🔍 分析 Word 模板: {template_path}")
    print("=" * 60)
    
    if not os.path.exists(template_path):
        print(f"❌ 檔案不存在: {template_path}")
        return
    
    try:
        # 載入 Word 文件
        print("📖 正在載入 Word 文件...")
        doc = Document(template_path)
        print("✅ Word 文件載入成功")
        
        print(f"📊 基本資訊:")
        print(f"   - 段落數: {len(doc.paragraphs)}")
        print(f"   - 表格數: {len(doc.tables)}")
        print(f"   - 區段數: {len(doc.sections)}")
        print()
        
        # 分析段落內容
        print("📝 段落內容分析:")
        print("-" * 40)
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"段落 {i+1}: {text}")
                
                # 檢查是否包含打勾相關文字
                if any(keyword in text for keyword in ['□', '✓', '財產保險', '保障需求', '車險']):
                    print(f"    ⭐ 包含關鍵字: {text}")
        
        print()
        
        # 分析表格內容
        print("📊 表格內容分析:")
        print("-" * 40)
        
        for table_idx, table in enumerate(doc.tables):
            print(f"表格 {table_idx + 1}:")
            print(f"  行數: {len(table.rows)}")
            print(f"  列數: {len(table.columns) if table.rows else 0}")
            
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(f"[{cell_idx+1}]:{cell_text}")
                
                if row_text:
                    print(f"    行 {row_idx + 1}: {' | '.join(row_text)}")
                    
                    # 檢查是否包含打勾相關文字
                    if any(keyword in ' '.join(row_text) for keyword in ['□', '✓', '財產保險', '保障需求', '車險']):
                        print(f"    ⭐ 包含關鍵字: {' | '.join(row_text)}")
            
            print()
        
        # 搜尋特定關鍵字
        print("🔍 關鍵字搜尋結果:")
        print("-" * 40)
        
        keywords = [
            '□財產保險', '✓財產保險',
            '□保障需求', '✓保障需求',
            '□否', '✓否',
            '□是', '✓是',
            '□車險', '✓車險',
            '□同上述之', '✓同上述之',
            '□商品保障內容符合客戶需求', '✓商品保障內容符合客戶需求'
        ]
        
        found_keywords = []
        
        # 在段落中搜尋
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            for keyword in keywords:
                if keyword in text:
                    found_keywords.append({
                        'type': 'paragraph',
                        'index': i + 1,
                        'keyword': keyword,
                        'context': text[:100] + '...' if len(text) > 100 else text
                    })
        
        # 在表格中搜尋
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text
                    for keyword in keywords:
                        if keyword in text:
                            found_keywords.append({
                                'type': 'table_cell',
                                'table': table_idx + 1,
                                'row': row_idx + 1,
                                'cell': cell_idx + 1,
                                'keyword': keyword,
                                'context': text[:100] + '...' if len(text) > 100 else text
                            })
        
        # 顯示搜尋結果
        if found_keywords:
            for item in found_keywords:
                if item['type'] == 'paragraph':
                    print(f"段落 {item['index']}: {item['keyword']}")
                    print(f"  內容: {item['context']}")
                else:
                    print(f"表格 {item['table']} 行 {item['row']} 格 {item['cell']}: {item['keyword']}")
                    print(f"  內容: {item['context']}")
                print()
        else:
            print("❌ 未找到指定的關鍵字")
        
        # 建議的填入策略
        print("💡 建議的填入策略:")
        print("-" * 40)
        
        if found_keywords:
            print("找到以下可填入的位置:")
            for item in found_keywords:
                if '財產保險' in item['keyword']:
                    location = item.get('index', f"表格{item.get('table')}行{item.get('row')}格{item.get('cell')}")
                    print(f"  - 財產保險選擇: {item['type']} {location}")
                elif '保障需求' in item['keyword']:
                    location = item.get('index', f"表格{item.get('table')}行{item.get('row')}格{item.get('cell')}")
                    print(f"  - 保障需求: {item['type']} {location}")
                elif '車險' in item['keyword']:
                    location = item.get('index', f"表格{item.get('table')}行{item.get('row')}格{item.get('cell')}")
                    print(f"  - 車險選擇: {item['type']} {location}")
                elif '否' in item['keyword']:
                    location = item.get('index', f"表格{item.get('table')}行{item.get('row')}格{item.get('cell')}")
                    print(f"  - 是否選擇: {item['type']} {location}")
        else:
            print("建議手動檢查 Word 檔案，找出需要填入的確切位置")
        
    except Exception as e:
        print(f"❌ 分析失敗: {str(e)}")

def main():
    """主函數"""
    template_path = "assets/templates/財產分析書.docx"
    analyze_word_template(template_path)

if __name__ == "__main__":
    main() 