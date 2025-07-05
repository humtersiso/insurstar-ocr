#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
檢查生成的 Word 檔案內容
"""

from docx import Document
import os

def check_generated_file(file_path):
    """檢查生成的 Word 檔案"""
    print(f"🔍 檢查檔案: {file_path}")
    print("=" * 50)
    
    if not os.path.exists(file_path):
        print(f"❌ 檔案不存在: {file_path}")
        return
    
    # 載入文檔
    doc = Document(file_path)
    
    # 檢查段落
    print(f"\n📝 段落內容:")
    print("-" * 30)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  段落 {i+1}: {para.text.strip()}")
    
    # 檢查表格
    print(f"\n📊 表格內容:")
    print("-" * 30)
    for table_idx, table in enumerate(doc.tables):
        print(f"  表格 {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"    行{row_idx+1} 格{cell_idx+1}: {cell.text.strip()}")
    
    # 檢查是否包含預期的內容
    print(f"\n🎯 檢查預期內容:")
    print("-" * 30)
    
    all_text = ""
    for para in doc.paragraphs:
        all_text += para.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + "\n"
    
    # 檢查關鍵內容
    expected_content = [
        "宏揚小客車租賃有限公司",  # 要保人
        "RAS-5879",              # 車牌號碼
        "PCN-RAS-5879-20250705", # PCN
        "新安東京海上產物保險股份有限公司",  # 保險公司
        "NT$27, 644"             # 保費
    ]
    
    for content in expected_content:
        if content in all_text:
            print(f"   ✅ {content}")
        else:
            print(f"   ❌ {content}")
    
    # 檢查是否還有未替換的標記
    print(f"\n🔍 檢查未替換的標記:")
    print("-" * 30)
    
    markers = ["{{PCN}}", "{{watermark_name_blue}}", "{{watermark_company_blue}}", 
               "{{policyholder}}", "{{license_number}}", "{{insurance_company}}"]
    
    for marker in markers:
        if marker in all_text:
            print(f"   ⚠️ {marker} - 未替換")
        else:
            print(f"   ✅ {marker} - 已替換或不存在")

def main():
    """主函數"""
    print("🔍 檢查生成的 Word 檔案")
    print("=" * 80)
    
    # 檢查最終測試檔案
    check_generated_file("test_outputs/final_test_simple_template.docx")
    
    print(f"\n📋 結論:")
    print(f"   如果看到預期內容（如要保人、車牌號碼等），表示標記替換成功")
    print(f"   如果看到未替換的標記，表示還有問題需要解決")

if __name__ == "__main__":
    main() 