#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
比較 word_template_processor_pure 和 test_pure_docxtpl 的差異
"""

import json
import os
from datetime import datetime
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from word_template_processor_pure import WordTemplateProcessorPure

def test_pure_docxtpl_method():
    """測試 test_pure_docxtpl 的方法"""
    print("🧪 測試 test_pure_docxtpl 方法")
    print("=" * 50)
    
    # 設定路徑
    TEMPLATE_PATH = 'assets/templates/財產分析書.docx'
    OUTPUT_PATH = 'test_outputs/method_pure_docxtpl.docx'
    OCR_JSON = 'ocr_results/gemini_ocr_output_20250704_144524.json'
    
    # 載入 OCR 資料
    with open(OCR_JSON, 'r', encoding='utf-8') as f:
        ocr_data = json.load(f)
    
    print(f"📄 OCR 資料載入成功")
    print(f"   - 保險公司: {ocr_data.get('insurance_company', 'N/A')}")
    print(f"   - 要保人: {ocr_data.get('policyholder', 'N/A')}")
    print(f"   - 車牌號碼: {ocr_data.get('license_number', 'N/A')}")
    
    # 載入模板
    tpl = DocxTemplate(TEMPLATE_PATH)
    
    # context: 先用 OCR 原始資料
    context = dict(ocr_data)
    
    # 補齊圖片欄位
    context['watermark_name_blue'] = InlineImage(tpl, 'assets/watermark/watermark_name_blue.png', width=Mm(30))
    context['watermark_company_blue'] = InlineImage(tpl, 'assets/watermark/watermark_company_blue.png', width=Mm(30))
    
    # 補其他欄位
    context['PCN'] = 'BB2H699299'
    
    print(f"\n📋 Context 資訊:")
    print(f"   - PCN: {context['PCN']}")
    print(f"   - watermark_name_blue: {type(context['watermark_name_blue'])}")
    print(f"   - watermark_company_blue: {type(context['watermark_company_blue'])}")
    print(f"   - Context 總欄位數: {len(context)}")
    
    # 渲染並儲存
    os.makedirs("test_outputs", exist_ok=True)
    tpl.render(context)
    tpl.save(OUTPUT_PATH)
    
    # 檢查檔案大小
    file_size = os.path.getsize(OUTPUT_PATH)
    print(f"\n✅ 檔案已生成: {OUTPUT_PATH}")
    print(f"📊 檔案大小: {file_size} bytes ({file_size/1024:.1f} KB)")
    
    return OUTPUT_PATH

def test_word_template_processor_pure():
    """測試 word_template_processor_pure 的方法"""
    print("\n🔧 測試 word_template_processor_pure 方法")
    print("=" * 50)
    
    # 載入 OCR 資料
    OCR_JSON = 'ocr_results/gemini_ocr_output_20250704_144524.json'
    with open(OCR_JSON, 'r', encoding='utf-8') as f:
        ocr_data = json.load(f)
    
    print(f"📄 OCR 資料載入成功")
    print(f"   - 保險公司: {ocr_data.get('insurance_company', 'N/A')}")
    print(f"   - 要保人: {ocr_data.get('policyholder', 'N/A')}")
    print(f"   - 車牌號碼: {ocr_data.get('license_number', 'N/A')}")
    
    # 建立處理器
    processor = WordTemplateProcessorPure()
    
    # 建立 context
    context = dict(ocr_data)
    
    # 補齊特殊欄位
    context['PCN'] = f"PCN-{ocr_data.get('license_number', '')}-{datetime.now().strftime('%Y%m%d')}"
    context['watermark_name_blue'] = 'assets/watermark/watermark_name_blue.png'
    context['watermark_company_blue'] = 'assets/watermark/watermark_company_blue.png'
    
    print(f"\n📋 Context 資訊:")
    print(f"   - PCN: {context['PCN']}")
    print(f"   - watermark_name_blue: {context['watermark_name_blue']}")
    print(f"   - watermark_company_blue: {context['watermark_company_blue']}")
    print(f"   - Context 總欄位數: {len(context)}")
    
    # 填入模板
    output_path = processor.fill_template(context, "test_outputs/method_processor_pure.docx")
    
    if output_path:
        # 檢查檔案大小
        file_size = os.path.getsize(output_path)
        print(f"\n✅ 檔案已生成: {output_path}")
        print(f"📊 檔案大小: {file_size} bytes ({file_size/1024:.1f} KB)")
        return output_path
    else:
        print(f"\n❌ 生成失敗")
        return None

def compare_files(file1, file2):
    """比較兩個檔案"""
    print(f"\n🔍 比較檔案")
    print("=" * 50)
    
    if not os.path.exists(file1) or not os.path.exists(file2):
        print(f"❌ 檔案不存在，無法比較")
        return
    
    # 比較檔案大小
    size1 = os.path.getsize(file1)
    size2 = os.path.getsize(file2)
    
    print(f"📊 檔案大小比較:")
    print(f"   {file1}: {size1} bytes ({size1/1024:.1f} KB)")
    print(f"   {file2}: {size2} bytes ({size2/1024:.1f} KB)")
    print(f"   差異: {abs(size1-size2)} bytes ({abs(size1-size2)/1024:.1f} KB)")
    
    # 檢查檔案內容
    from docx import Document
    
    def get_file_content(file_path):
        doc = Document(file_path)
        content = ""
        for para in doc.paragraphs:
            content += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + "\n"
        return content
    
    content1 = get_file_content(file1)
    content2 = get_file_content(file2)
    
    # 檢查關鍵內容
    key_content = [
        "宏揚小客車租賃有限公司",  # 要保人
        "RAS-5879",              # 車牌號碼
        "新安東京海上產物保險股份有限公司",  # 保險公司
        "NT$27, 644"             # 保費
    ]
    
    print(f"\n🎯 關鍵內容檢查:")
    for content in key_content:
        in_file1 = content in content1
        in_file2 = content in content2
        print(f"   {content}:")
        print(f"     {file1}: {'✅' if in_file1 else '❌'}")
        print(f"     {file2}: {'✅' if in_file2 else '❌'}")
    
    # 檢查 PCN
    pcn1 = "BB2H699299" in content1
    pcn2 = "PCN-RAS-5879" in content2
    print(f"\n🔢 PCN 檢查:")
    print(f"   {file1}: BB2H699299 {'✅' if pcn1 else '❌'}")
    print(f"   {file2}: PCN-RAS-5879 {'✅' if pcn2 else '❌'}")

def main():
    """主函數"""
    print("🔍 比較兩種方法")
    print("=" * 80)
    
    # 測試兩種方法
    file1 = test_pure_docxtpl_method()
    file2 = test_word_template_processor_pure()
    
    # 比較結果
    if file1 and file2:
        compare_files(file1, file2)
    
    print(f"\n📋 總結:")
    print(f"   1. test_pure_docxtpl: 直接使用 DocxTemplate，簡單直接")
    print(f"   2. word_template_processor_pure: 使用處理器類別，功能更完整")
    print(f"   3. 主要差異可能在 context 建立方式和圖片處理方式")

if __name__ == "__main__":
    main() 