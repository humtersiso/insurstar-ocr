#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
建立簡單的 Word 模板
"""

from docx import Document
from docx.shared import Inches
import os

def create_simple_template():
    """建立簡單的 Word 模板"""
    print("📝 建立簡單 Word 模板")
    
    # 建立新文檔
    doc = Document()
    
    # 添加標題
    title = doc.add_heading('財產分析書', 0)
    
    # 添加表格
    table = doc.add_table(rows=25, cols=15)
    table.style = 'Table Grid'
    
    # 設定表格標題行
    header_row = table.rows[0]
    header_cells = header_row.cells
    header_cells[0].text = '保險類型'
    header_cells[1].text = '要保人'
    header_cells[2].text = '被保險人'
    header_cells[3].text = '車牌號碼'
    header_cells[4].text = '保費'
    
    # 在表格中插入標記
    # 第3行：要保人資訊
    row3 = table.rows[2]
    row3.cells[1].text = '{{policyholder}}'
    row3.cells[2].text = '{{insured_person}}'
    row3.cells[3].text = '{{license_number}}'
    row3.cells[4].text = '{{total_premium}}'
    
    # 第4行：身分證號
    row4 = table.rows[3]
    row4.cells[1].text = '{{policyholder_id}}'
    row4.cells[2].text = '{{id_number}}'
    
    # 第5行：出生日期
    row5 = table.rows[4]
    row5.cells[1].text = '{{policyholder_birth_date}}'
    row5.cells[2].text = '{{birth_date}}'
    
    # 第6行：法定代理人
    row6 = table.rows[5]
    row6.cells[1].text = '{{policyholder_legal_representative}}'
    row6.cells[2].text = '{{legal_representative}}'
    
    # 第22行：浮水印和 PCN
    row22 = table.rows[21]
    row22.cells[1].text = '{{watermark_name_blue}}'
    row22.cells[2].text = '{{watermark_name_blue}}'
    row22.cells[3].text = '{{watermark_name_blue}}'
    row22.cells[5].text = '{{PCN}}'
    row22.cells[6].text = '{{PCN}}'
    row22.cells[7].text = '{{PCN}}'
    row22.cells[10].text = '{{watermark_company_blue}}'
    row22.cells[11].text = '{{watermark_company_blue}}'
    row22.cells[12].text = '{{watermark_company_blue}}'
    
    # 添加一些說明文字
    doc.add_paragraph('保險公司：{{insurance_company}}')
    doc.add_paragraph('強制險期間：{{compulsory_insurance_period}}')
    doc.add_paragraph('任意險期間：{{optional_insurance_period}}')
    
    # 儲存模板
    output_path = "assets/templates/財產分析書_simple.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    print(f"✅ 簡單模板已建立: {output_path}")
    return output_path

def test_simple_template():
    """測試簡單模板"""
    print("\n🧪 測試簡單模板")
    print("-" * 40)
    
    import json
    from datetime import datetime
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm
    
    # 載入 OCR 資料
    with open('ocr_results/gemini_ocr_output_20250704_144524.json', 'r', encoding='utf-8') as f:
        ocr_data = json.load(f)
    
    # 載入模板
    template_path = "assets/templates/財產分析書_simple.docx"
    doc = DocxTemplate(template_path)
    
    # 建立 context
    context = dict(ocr_data)
    context['PCN'] = f"PCN-{ocr_data.get('license_number', '')}-{datetime.now().strftime('%Y%m%d')}"
    context['watermark_name_blue'] = InlineImage(doc, 'assets/watermark/watermark_name_blue.png', width=Mm(30))
    context['watermark_company_blue'] = InlineImage(doc, 'assets/watermark/watermark_company_blue.png', width=Mm(30))
    
    # 檢查模板中的標記
    print(f"🔍 檢查模板標記:")
    all_text = ""
    if doc.docx and doc.docx.paragraphs:
        for para in doc.docx.paragraphs:
            all_text += para.text + "\n"
    
    if doc.docx and doc.docx.tables:
        for table in doc.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text += para.text + "\n"
    
    # 檢查關鍵標記
    key_markers = ['{{PCN}}', '{{watermark_name_blue}}', '{{watermark_company_blue}}', '{{policyholder}}', '{{license_number}}']
    found_markers = []
    for marker in key_markers:
        if marker in all_text:
            print(f"   ✅ {marker} - 存在於模板中")
            found_markers.append(marker)
        else:
            print(f"   ❌ {marker} - 不存在於模板中")
    
    # 渲染並儲存
    output_path = "test_outputs/test_simple_template.docx"
    os.makedirs("test_outputs", exist_ok=True)
    
    try:
        doc.render(context)
        doc.save(output_path)
        print(f"\n✅ 檔案已生成: {output_path}")
        
        # 檢查檔案大小
        file_size = os.path.getsize(output_path)
        print(f"📊 檔案大小: {file_size} bytes ({file_size/1024:.1f} KB)")
        
        if len(found_markers) > 0:
            print(f"🎯 成功找到 {len(found_markers)} 個標記")
        else:
            print(f"⚠️ 沒有找到任何標記")
            
    except Exception as e:
        print(f"❌ 生成失敗: {str(e)}")

def main():
    """主函數"""
    print("📝 建立簡單 Word 模板")
    print("=" * 50)
    
    # 建立簡單模板
    template_path = create_simple_template()
    
    # 測試模板
    test_simple_template()
    
    print(f"\n📁 結果:")
    print(f"   1. 模板檔案: {template_path}")
    print(f"   2. 測試檔案: test_outputs/test_simple_template.docx")

if __name__ == "__main__":
    main() 