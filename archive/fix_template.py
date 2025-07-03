#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修復Word模板中的語法錯誤
"""

import os
from docx import Document
import re

def fix_template_syntax():
    """修復模板中的語法錯誤"""
    print("🔧 修復Word模板語法錯誤")
    print("=" * 50)
    
    template_path = "assets/templates/財產分析書.docx"
    fixed_path = "assets/templates/財產分析書_fixed.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ 模板檔案不存在: {template_path}")
        return False
    
    try:
        # 載入原始模板
        doc = Document(template_path)
        
        # 修復表格中的語法錯誤
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # 修復空的模板標記
                        if "{{}}" in paragraph.text:
                            print(f"🔧 修復空標記: {paragraph.text}")
                            paragraph.text = paragraph.text.replace("{{}}", "{{gender}}")
                        
                        # 修復其他可能的語法問題
                        # 檢查是否有未閉合的標記
                        text = paragraph.text
                        open_count = text.count("{{")
                        close_count = text.count("}}")
                        
                        if open_count != close_count:
                            print(f"⚠️ 標記不匹配: {text}")
                            # 移除有問題的標記
                            text = re.sub(r'\{\{[^}]*$', '', text)
                            text = re.sub(r'^\}[^}]*\}\}', '', text)
                            paragraph.text = text
        
        # 修復段落中的語法錯誤
        for paragraph in doc.paragraphs:
            if "{{}}" in paragraph.text:
                print(f"🔧 修復段落空標記: {paragraph.text}")
                paragraph.text = paragraph.text.replace("{{}}", "{{gender}}")
        
        # 儲存修復後的模板
        doc.save(fixed_path)
        print(f"✅ 修復完成: {fixed_path}")
        
        return True
        
    except Exception as e:
        print(f"❌ 修復失敗: {str(e)}")
        return False

def create_simple_template():
    """創建一個簡化的測試模板"""
    print("\n📝 創建簡化測試模板")
    print("=" * 50)
    
    doc = Document()
    
    # 添加標題
    title = doc.add_heading('財產分析書', 0)
    
    # 添加基本資訊表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 表格標題
    table.cell(0, 0).text = "欄位"
    table.cell(0, 1).text = "內容"
    
    # 基本欄位
    fields = [
        ("保險公司", "{{insurance_company}}"),
        ("要保人", "{{policyholder}}"),
        ("被保險人", "{{insured_person}}"),
        ("車牌號碼", "{{license_number}}")
    ]
    
    for i, (field, value) in enumerate(fields, 1):
        table.cell(i, 0).text = field
        table.cell(i, 1).text = value
    
    # 添加勾選項目
    doc.add_heading('保險類型', level=1)
    
    # 勾選項目段落
    p1 = doc.add_paragraph('□人身保險 ')
    p1.add_run('{{CHECK_1}}').bold = True
    p1.add_run(' 財產保險 □ 旅行平安保險')
    
    p2 = doc.add_paragraph('□強制險 ')
    p2.add_run('{{CHECK_2}}').bold = True
    p2.add_run(' 任意車險')
    
    # 添加性別選擇
    doc.add_heading('基本資料', level=1)
    gender_p = doc.add_paragraph('性別: □男 ')
    gender_p.add_run('{{gender_male}}').bold = True
    gender_p.add_run(' □女 ')
    gender_p.add_run('{{gender_female}}').bold = True
    
    # 儲存簡化模板
    simple_path = "assets/templates/簡化財產分析書模板.docx"
    doc.save(simple_path)
    print(f"✅ 簡化模板已創建: {simple_path}")
    
    return simple_path

def main():
    """主函數"""
    print("🔧 Word模板修復工具")
    print("=" * 80)
    
    # 修復原始模板
    fix_success = fix_template_syntax()
    
    # 創建簡化模板
    simple_path = create_simple_template()
    
    print("\n📊 修復結果:")
    print(f"原始模板修復: {'✅ 成功' if fix_success else '❌ 失敗'}")
    print(f"簡化模板創建: ✅ 成功")
    
    if fix_success:
        print("\n🎉 模板修復完成！")
        print("📝 建議使用修復後的模板進行測試")
    else:
        print("\n⚠️ 使用簡化模板進行測試")

if __name__ == "__main__":
    main() 