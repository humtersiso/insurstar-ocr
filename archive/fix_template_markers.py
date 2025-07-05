#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修復 Word 模板中被拆分的標記
"""

import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def fix_split_markers(template_path, output_path):
    """修復被拆分的標記"""
    print(f"🔧 修復模板: {template_path}")
    print(f"📁 輸出到: {output_path}")
    
    # 載入文檔
    doc = Document(template_path)
    
    # 要修復的標記列表
    target_markers = [
        '{{PCN}}',
        '{{watermark_name_blue}}',
        '{{watermark_company_blue}}',
        '{{policyholder}}',
        '{{license_number}}',
        '{{total_premium}}',
        '{{insured_person}}',
        '{{legal_representative}}',
        '{{id_number}}',
        '{{birth_date}}',
        '{{policyholder_legal_representative}}',
        '{{policyholder_id}}',
        '{{policyholder_birth_date}}',
        '{{vehicle_type}}',
        '{{insurance_company}}',
        '{{compulsory_insurance_period}}',
        '{{optional_insurance_period}}'
    ]
    
    fixed_count = 0
    
    # 修復段落中的標記
    for paragraph in doc.paragraphs:
        fixed_count += fix_paragraph_markers(paragraph, target_markers)
    
    # 修復表格中的標記
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fixed_count += fix_paragraph_markers(paragraph, target_markers)
    
    # 儲存修復後的文檔
    doc.save(output_path)
    
    print(f"✅ 修復完成！共修復 {fixed_count} 個標記")
    return fixed_count

def fix_paragraph_markers(paragraph, target_markers):
    """修復段落中被拆分的標記"""
    fixed_count = 0
    
    # 獲取段落的所有文字運行
    runs = list(paragraph.runs)
    if len(runs) <= 1:
        return 0
    
    # 檢查每個標記
    for marker in target_markers:
        # 檢查標記是否被拆分
        if is_marker_split(runs, marker):
            print(f"   🔧 修復標記: {marker}")
            if fix_split_marker_in_runs(runs, marker):
                fixed_count += 1
    
    return fixed_count

def is_marker_split(runs, marker):
    """檢查標記是否被拆分"""
    # 將所有運行的文字連接起來
    full_text = ''.join([run.text for run in runs])
    return marker in full_text

def fix_split_marker_in_runs(runs, marker):
    """修復被拆分的標記"""
    # 找到包含標記的運行
    marker_runs = []
    current_text = ""
    
    for i, run in enumerate(runs):
        current_text += run.text
        if marker in current_text:
            # 找到標記的開始位置
            start_pos = current_text.find(marker)
            end_pos = start_pos + len(marker)
            
            # 計算標記在每個運行中的位置
            temp_text = ""
            for j in range(i + 1):
                if j == i:
                    # 當前運行
                    run_start = len(temp_text)
                    run_end = len(temp_text) + len(run.text)
                    
                    if start_pos < run_end and end_pos > run_start:
                        marker_runs.append((j, run))
                temp_text += runs[j].text
            
            break
    
    if not marker_runs:
        return False
    
    # 如果標記跨越多個運行，需要合併
    if len(marker_runs) > 1:
        # 合併所有相關運行
        first_run = marker_runs[0][1]
        first_run.text = marker
        
        # 清除其他運行
        for _, run in marker_runs[1:]:
            run.text = ""
        
        return True
    elif len(marker_runs) == 1:
        # 標記在單個運行中，但可能被其他文字包圍
        run = marker_runs[0][1]
        if run.text != marker:
            run.text = marker
            return True
    
    return False

def create_clean_template():
    """建立一個乾淨的模板，手動插入標記"""
    print("🧹 建立乾淨模板")
    
    # 載入原始模板
    doc = Document("assets/templates/財產分析書.docx")
    
    # 在適當位置插入標記
    # 這裡我們在表格的第22行插入標記
    if len(doc.tables) > 0:
        table = doc.tables[0]
        if len(table.rows) >= 22:
            row = table.rows[21]  # 第22行 (0-based)
            
            # 插入標記到特定格子
            if len(row.cells) >= 8:
                # 第6格：PCN
                row.cells[5].text = "{{PCN}}"
                
                # 第2-4格：watermark_name_blue
                for i in range(1, 4):
                    row.cells[i].text = "{{watermark_name_blue}}"
                
                # 第11-13格：watermark_company_blue
                for i in range(10, 13):
                    row.cells[i].text = "{{watermark_company_blue}}"
    
    # 儲存乾淨模板
    output_path = "assets/templates/財產分析書_clean.docx"
    doc.save(output_path)
    print(f"✅ 乾淨模板已建立: {output_path}")
    return output_path

def main():
    """主函數"""
    print("🔧 Word 模板標記修復工具")
    print("=" * 50)
    
    # 方法1：嘗試修復現有模板
    template_path = "assets/templates/財產分析書.docx"
    output_path = "assets/templates/財產分析書_fixed_markers.docx"
    
    if os.path.exists(template_path):
        fixed_count = fix_split_markers(template_path, output_path)
        if fixed_count > 0:
            print(f"✅ 修復成功！共修復 {fixed_count} 個標記")
        else:
            print("⚠️ 沒有找到需要修復的標記")
    else:
        print(f"❌ 模板檔案不存在: {template_path}")
    
    # 方法2：建立乾淨模板
    print(f"\n🧹 建立乾淨模板...")
    clean_template_path = create_clean_template()
    
    print(f"\n📁 修復結果:")
    print(f"   1. 修復模板: {output_path}")
    print(f"   2. 乾淨模板: {clean_template_path}")
    print(f"\n💡 建議使用乾淨模板進行測試")

if __name__ == "__main__":
    main() 