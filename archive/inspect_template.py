#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
檢查Word模板檔案內容
"""

import os
from docx import Document
import re

def inspect_template_content():
    """檢查模板檔案內容"""
    print("🔍 檢查Word模板檔案內容")
    print("=" * 60)
    
    template_path = "assets/templates/財產分析書.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ 模板檔案不存在: {template_path}")
        return
    
    try:
        doc = Document(template_path)
        
        print(f"📄 模板檔案: {template_path}")
        print(f"📊 基本資訊:")
        print(f"   - 段落數量: {len(doc.paragraphs)}")
        print(f"   - 表格數量: {len(doc.tables)}")
        print(f"   - 區段數量: {len(doc.sections)}")
        
        # 檢查所有段落內容
        print(f"\n📝 段落內容檢查:")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"   段落 {i+1}: {text}")
                
                # 檢查是否有問題的標記
                if '{{' in text and '}}' in text:
                    print(f"     -> 包含模板標記")
                
                # 檢查是否有語法錯誤的標記
                if '{{' in text and '}}' not in text:
                    print(f"     -> ⚠️ 可能有未閉合的標記")
                if '}}' in text and '{{' not in text:
                    print(f"     -> ⚠️ 可能有未開頭的標記")
        
        # 檢查表格內容
        if doc.tables:
            print(f"\n📊 表格內容檢查:")
            for table_idx, table in enumerate(doc.tables):
                print(f"   表格 {table_idx + 1}:")
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            text = para.text.strip()
                            if text:
                                print(f"      [{row_idx},{cell_idx}] 段落 {para_idx}: {text}")
                                
                                # 檢查是否有問題的標記
                                if '{{' in text and '}}' in text:
                                    print(f"        -> 包含模板標記")
                                
                                # 檢查是否有語法錯誤的標記
                                if '{{' in text and '}}' not in text:
                                    print(f"        -> ⚠️ 可能有未閉合的標記")
                                if '}}' in text and '{{' not in text:
                                    print(f"        -> ⚠️ 可能有未開頭的標記")
        
        # 搜尋所有可能的模板標記
        print(f"\n🔍 搜尋所有模板標記:")
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text += para.text + "\n"
        
        # 使用正則表達式找出所有模板標記
        template_pattern = r'\{\{[^}]*\}\}'
        matches = re.findall(template_pattern, all_text)
        
        if matches:
            print(f"   找到 {len(matches)} 個模板標記:")
            for match in matches:
                print(f"     - {match}")
        else:
            print("   沒有找到標準的模板標記")
        
        # 檢查是否有語法錯誤的標記
        open_braces = all_text.count('{{')
        close_braces = all_text.count('}}')
        
        print(f"\n🔍 標記語法檢查:")
        print(f"   開頭標記 '{{{{': {open_braces}")
        print(f"   結尾標記 '}}': {close_braces}")
        
        if open_braces != close_braces:
            print(f"   ⚠️ 標記數量不匹配！")
        else:
            print(f"   ✅ 標記數量匹配")
        
        # 搜尋可能的問題標記
        print(f"\n🔍 搜尋可能的問題標記:")
        problem_patterns = [
            r'\{\{[^}]*$',  # 未閉合的標記
            r'^[^}]*\}\}',  # 未開頭的標記
            r'\{\{[^}]*\n[^}]*\}\}',  # 跨行的標記
        ]
        
        for pattern in problem_patterns:
            matches = re.findall(pattern, all_text, re.MULTILINE)
            if matches:
                print(f"   找到問題標記 (模式: {pattern}):")
                for match in matches:
                    print(f"     - {repr(match)}")
        
    except Exception as e:
        print(f"❌ 檢查失敗: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    inspect_template_content() 