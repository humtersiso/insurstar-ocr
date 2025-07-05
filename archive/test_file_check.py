#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
檢查 Word 檔案內容
"""

from docx import Document
import os

def check_word_file(file_path):
    """檢查 Word 檔案"""
    print(f"🔍 檢查檔案: {file_path}")
    print("=" * 50)
    
    if not os.path.exists(file_path):
        print(f"❌ 檔案不存在: {file_path}")
        return
    
    try:
        doc = Document(file_path)
        
        print(f"📊 基本資訊:")
        print(f"   - 段落數量: {len(doc.paragraphs)}")
        print(f"   - 表格數量: {len(doc.tables)}")
        
        print(f"\n📝 段落內容:")
        print("-" * 30)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"  段落 {i+1}: {para.text.strip()}")
        
        print(f"\n📊 表格內容:")
        print("-" * 30)
        for table_idx, table in enumerate(doc.tables):
            print(f"  表格 {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        print(f"    行{row_idx+1} 格{cell_idx+1}: {cell.text.strip()}")
        
        # 檢查字體設定
        print(f"\n🔤 字體檢查:")
        print("-" * 30)
        font_info = {}
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    font_name = run.font.name if run.font.name else "預設"
                    font_size = run.font.size if run.font.size else "預設"
                    font_info[f"{font_name}_{font_size}"] = font_info.get(f"{font_name}_{font_size}", 0) + 1
        
        for font, count in font_info.items():
            print(f"  {font}: {count} 個文字運行")
        
        print(f"\n✅ 檢查完成")
        
    except Exception as e:
        print(f"❌ 檢查失敗: {str(e)}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        # 使用命令列參數指定的檔案
        check_word_file(sys.argv[1])
    else:
        # 檢查最新的檔案
        check_word_file("property_reports/財產分析書_20250705_071043.docx") 