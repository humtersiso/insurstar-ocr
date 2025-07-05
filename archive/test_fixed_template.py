#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
測試修復後的模板
"""

import json
import os
from datetime import datetime
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm

def test_fixed_template():
    """測試修復後的模板"""
    print("🧪 測試修復後的模板")
    print("=" * 50)
    
    # 載入 OCR 資料
    with open('ocr_results/gemini_ocr_output_20250704_144524.json', 'r', encoding='utf-8') as f:
        ocr_data = json.load(f)
    
    print(f"📄 OCR 資料載入成功")
    print(f"   - 保險公司: {ocr_data.get('insurance_company', 'N/A')}")
    print(f"   - 要保人: {ocr_data.get('policyholder', 'N/A')}")
    print(f"   - 車牌號碼: {ocr_data.get('license_number', 'N/A')}")
    
    # 測試修復後的模板
    template_paths = [
        "assets/templates/財產分析書_fixed_markers.docx",
        "assets/templates/財產分析書_clean.docx"
    ]
    
    for i, template_path in enumerate(template_paths, 1):
        if not os.path.exists(template_path):
            print(f"❌ 模板不存在: {template_path}")
            continue
            
        print(f"\n🔧 測試模板 {i}: {template_path}")
        print("-" * 40)
        
        # 載入模板
        doc = DocxTemplate(template_path)
        
        # 建立 context
        context = dict(ocr_data)
        context['PCN'] = f"PCN-{ocr_data.get('license_number', '')}-{datetime.now().strftime('%Y%m%d')}"
        context['watermark_name_blue'] = InlineImage(doc, 'assets/watermark/watermark_name_blue.png', width=Mm(30))
        context['watermark_company_blue'] = InlineImage(doc, 'assets/watermark/watermark_company_blue.png', width=Mm(30))
        
        # 檢查模板中的標記
        print(f"🔍 檢查模板標記:")
        all_text = ""
        if doc.docx and doc.docx.paragraphs:
            for para in doc.docx.paragraphs:
                all_text += para.text + "\n"
        
        if doc.docx and doc.docx.tables:
            for table in doc.docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            all_text += para.text + "\n"
        
        # 檢查關鍵標記
        key_markers = ['{{PCN}}', '{{watermark_name_blue}}', '{{watermark_company_blue}}', '{{policyholder}}', '{{license_number}}']
        found_markers = []
        for marker in key_markers:
            if marker in all_text:
                print(f"   ✅ {marker} - 存在於模板中")
                found_markers.append(marker)
            else:
                print(f"   ❌ {marker} - 不存在於模板中")
        
        # 渲染並儲存
        output_path = f"test_outputs/test_fixed_template_{i}.docx"
        os.makedirs("test_outputs", exist_ok=True)
        
        try:
            doc.render(context)
            doc.save(output_path)
            print(f"\n✅ 檔案已生成: {output_path}")
            
            # 檢查檔案大小
            file_size = os.path.getsize(output_path)
            print(f"📊 檔案大小: {file_size} bytes ({file_size/1024:.1f} KB)")
            
            if len(found_markers) > 0:
                print(f"🎯 成功找到 {len(found_markers)} 個標記")
            else:
                print(f"⚠️ 沒有找到任何標記")
                
        except Exception as e:
            print(f"❌ 生成失敗: {str(e)}")

def main():
    """主函數"""
    print("🧪 修復模板測試")
    print("=" * 80)
    
    test_fixed_template()
    
    print(f"\n📁 請檢查 test_outputs/ 目錄下的檔案")
    print(f"   1. test_fixed_template_1.docx - 修復標記模板測試")
    print(f"   2. test_fixed_template_2.docx - 乾淨模板測試")

if __name__ == "__main__":
    main() 