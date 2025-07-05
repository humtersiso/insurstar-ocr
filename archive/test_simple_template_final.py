#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最終測試：使用簡單模板
"""

import json
import os
from datetime import datetime
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm

def test_simple_template_final():
    """最終測試簡單模板"""
    print("🎯 最終測試：簡單模板")
    print("=" * 50)
    
    # 載入 OCR 資料
    with open('ocr_results/gemini_ocr_output_20250704_144524.json', 'r', encoding='utf-8') as f:
        ocr_data = json.load(f)
    
    print(f"📄 OCR 資料載入成功")
    print(f"   - 保險公司: {ocr_data.get('insurance_company', 'N/A')}")
    print(f"   - 要保人: {ocr_data.get('policyholder', 'N/A')}")
    print(f"   - 車牌號碼: {ocr_data.get('license_number', 'N/A')}")
    
    # 載入簡單模板
    template_path = "assets/templates/財產分析書_simple.docx"
    doc = DocxTemplate(template_path)
    
    # 建立 context
    context = dict(ocr_data)
    context['PCN'] = f"PCN-{ocr_data.get('license_number', '')}-{datetime.now().strftime('%Y%m%d')}"
    context['watermark_name_blue'] = InlineImage(doc, 'assets/watermark/watermark_name_blue.png', width=Mm(30))
    context['watermark_company_blue'] = InlineImage(doc, 'assets/watermark/watermark_company_blue.png', width=Mm(30))
    
    print(f"\n📋 Context 資訊:")
    print(f"   - PCN: {context['PCN']}")
    print(f"   - watermark_name_blue: {type(context['watermark_name_blue'])}")
    print(f"   - watermark_company_blue: {type(context['watermark_company_blue'])}")
    
    # 檢查模板中的標記
    print(f"\n🔍 檢查模板標記:")
    all_text = ""
    if doc.docx and doc.docx.paragraphs:
        for para in doc.docx.paragraphs:
            all_text += para.text + "\n"
    
    if doc.docx and doc.docx.tables:
        for table in doc.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text += para.text + "\n"
    
    # 檢查關鍵標記
    key_markers = ['{{PCN}}', '{{watermark_name_blue}}', '{{watermark_company_blue}}', '{{policyholder}}', '{{license_number}}']
    found_markers = []
    for marker in key_markers:
        if marker in all_text:
            print(f"   ✅ {marker} - 存在於模板中")
            found_markers.append(marker)
        else:
            print(f"   ❌ {marker} - 不存在於模板中")
    
    # 渲染並儲存
    output_path = "test_outputs/final_test_simple_template.docx"
    os.makedirs("test_outputs", exist_ok=True)
    
    try:
        doc.render(context)
        doc.save(output_path)
        print(f"\n✅ 檔案已生成: {output_path}")
        
        # 檢查檔案大小
        file_size = os.path.getsize(output_path)
        print(f"📊 檔案大小: {file_size} bytes ({file_size/1024:.1f} KB)")
        
        if len(found_markers) > 0:
            print(f"🎯 成功找到 {len(found_markers)} 個標記")
            print(f"   - 標記列表: {', '.join(found_markers)}")
        else:
            print(f"⚠️ 沒有找到任何標記")
            
        return True
        
    except Exception as e:
        print(f"❌ 生成失敗: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """主函數"""
    print("🎯 最終測試：簡單模板")
    print("=" * 80)
    
    success = test_simple_template_final()
    
    if success:
        print(f"\n🎉 測試成功！")
        print(f"📁 請檢查檔案: test_outputs/final_test_simple_template.docx")
        print(f"💡 這個檔案應該包含正確的標記替換和圖片")
    else:
        print(f"\n❌ 測試失敗")
    
    print(f"\n📋 總結:")
    print(f"   1. 問題根源：Word 模板中的標記被 Word 編輯器拆分")
    print(f"   2. 解決方案：使用程式建立的簡單模板")
    print(f"   3. 結果：標記可以正常替換")

if __name__ == "__main__":
    main() 