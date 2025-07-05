#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
簡單驗證腳本：檢查標記替換是否有效
"""

import json
import os
from datetime import datetime
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm

def test_simple_verification():
    """簡單驗證測試"""
    print("🔍 簡單驗證測試")
    print("=" * 50)
    
    # 載入 OCR 資料
    with open('ocr_results/gemini_ocr_output_20250704_144524.json', 'r', encoding='utf-8') as f:
        ocr_data = json.load(f)
    
    print(f"📄 OCR 資料載入成功")
    print(f"   - 保險公司: {ocr_data.get('insurance_company', 'N/A')}")
    print(f"   - 要保人: {ocr_data.get('policyholder', 'N/A')}")
    print(f"   - 車牌號碼: {ocr_data.get('license_number', 'N/A')}")
    
    # 載入模板
    template_path = "assets/templates/財產分析書.docx"
    doc = DocxTemplate(template_path)
    
    # 建立 context（像 test_pure_docxtpl.py 一樣）
    context = dict(ocr_data)
    
    # 補齊特殊欄位
    context['PCN'] = f"PCN-{ocr_data.get('license_number', '')}-{datetime.now().strftime('%Y%m%d')}"
    context['watermark_name_blue'] = InlineImage(doc, 'assets/watermark/watermark_name_blue.png', width=Mm(30))
    context['watermark_company_blue'] = InlineImage(doc, 'assets/watermark/watermark_company_blue.png', width=Mm(30))
    
    print(f"\n📋 Context 資訊:")
    print(f"   - PCN: {context['PCN']}")
    print(f"   - watermark_name_blue: {type(context['watermark_name_blue'])}")
    print(f"   - watermark_company_blue: {type(context['watermark_company_blue'])}")
    
    # 檢查模板中的標記
    print(f"\n🔍 檢查模板標記:")
    all_text = ""
    if doc.docx and doc.docx.paragraphs:
        for para in doc.docx.paragraphs:
            all_text += para.text + "\n"
    
    if doc.docx and doc.docx.tables:
        for table in doc.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text += para.text + "\n"
    
    # 檢查關鍵標記
    key_markers = ['{{PCN}}', '{{watermark_name_blue}}', '{{watermark_company_blue}}', '{{policyholder}}', '{{license_number}}']
    for marker in key_markers:
        if marker in all_text:
            print(f"   ✅ {marker} - 存在於模板中")
        else:
            print(f"   ❌ {marker} - 不存在於模板中")
    
    # 渲染並儲存
    output_path = "test_outputs/simple_verification.docx"
    os.makedirs("test_outputs", exist_ok=True)
    
    try:
        doc.render(context)
        doc.save(output_path)
        print(f"\n✅ 檔案已生成: {output_path}")
        
        # 檢查檔案大小
        file_size = os.path.getsize(output_path)
        print(f"📊 檔案大小: {file_size} bytes ({file_size/1024:.1f} KB)")
        
        return True
        
    except Exception as e:
        print(f"❌ 生成失敗: {str(e)}")
        return False

def test_with_fixed_template():
    """使用修復後的模板測試"""
    print(f"\n🔧 使用修復模板測試")
    print("-" * 50)
    
    # 載入 OCR 資料
    with open('ocr_results/gemini_ocr_output_20250704_144524.json', 'r', encoding='utf-8') as f:
        ocr_data = json.load(f)
    
    # 載入修復模板
    template_path = "assets/templates/財產分析書_fixed.docx"
    doc = DocxTemplate(template_path)
    
    # 建立 context
    context = dict(ocr_data)
    context['PCN'] = f"PCN-{ocr_data.get('license_number', '')}-{datetime.now().strftime('%Y%m%d')}"
    context['watermark_name_blue'] = InlineImage(doc, 'assets/watermark/watermark_name_blue.png', width=Mm(30))
    context['watermark_company_blue'] = InlineImage(doc, 'assets/watermark/watermark_company_blue.png', width=Mm(30))
    
    # 渲染並儲存
    output_path = "test_outputs/simple_verification_fixed.docx"
    
    try:
        doc.render(context)
        doc.save(output_path)
        print(f"✅ 修復模板檔案已生成: {output_path}")
        
        # 檢查檔案大小
        file_size = os.path.getsize(output_path)
        print(f"📊 檔案大小: {file_size} bytes ({file_size/1024:.1f} KB)")
        
        return True
        
    except Exception as e:
        print(f"❌ 生成失敗: {str(e)}")
        return False

def main():
    """主函數"""
    print("🧪 簡單驗證測試")
    print("=" * 80)
    
    # 測試原始模板
    test_simple_verification()
    
    # 測試修復模板
    test_with_fixed_template()
    
    print(f"\n📁 請檢查 test_outputs/ 目錄下的檔案")
    print(f"   1. simple_verification.docx - 原始模板測試")
    print(f"   2. simple_verification_fixed.docx - 修復模板測試")

if __name__ == "__main__":
    main() 