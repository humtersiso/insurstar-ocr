#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
測試Word模板基本功能
"""

import os
from docx import Document
from docxtpl import DocxTemplate

def test_template_basic():
    """測試模板基本功能"""
    print("🔍 測試Word模板基本功能")
    print("=" * 50)
    
    template_path = "assets/templates/財產分析書.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ 模板檔案不存在: {template_path}")
        return False
    
    try:
        # 使用python-docx讀取模板
        print("📄 使用python-docx讀取模板...")
        doc = Document(template_path)
        print(f"✅ 模板載入成功")
        print(f"   - 段落數量: {len(doc.paragraphs)}")
        print(f"   - 表格數量: {len(doc.tables)}")
        print(f"   - 區段數量: {len(doc.sections)}")
        
        # 檢查前幾個段落的內容
        print("\n📝 前5個段落內容:")
        for i, para in enumerate(doc.paragraphs[:5]):
            text = para.text.strip()
            if text:
                print(f"   {i+1}. {text[:100]}...")
        
        # 檢查表格內容
        if doc.tables:
            print(f"\n📊 表格資訊:")
            for i, table in enumerate(doc.tables):
                print(f"   表格 {i+1}: {len(table.rows)} 行 x {len(table.columns)} 列")
        
        return True
        
    except Exception as e:
        print(f"❌ 讀取模板失敗: {str(e)}")
        return False

def test_docxtpl_basic():
    """測試docxtpl基本功能"""
    print("\n🔍 測試docxtpl基本功能")
    print("=" * 50)
    
    template_path = "assets/templates/財產分析書.docx"
    
    try:
        # 使用docxtpl讀取模板
        print("📄 使用docxtpl讀取模板...")
        template = DocxTemplate(template_path)
        print(f"✅ docxtpl模板載入成功")
        
        # 檢查模板變數
        print("🔍 檢查模板變數...")
        variables = template.get_undeclared_template_variables()
        print(f"   找到 {len(variables)} 個模板變數:")
        for var in variables:
            print(f"   - {var}")
        
        return True
        
    except Exception as e:
        print(f"❌ docxtpl讀取失敗: {str(e)}")
        return False

def test_simple_render():
    """測試簡單的模板渲染"""
    print("\n🔍 測試簡單模板渲染")
    print("=" * 50)
    
    template_path = "assets/templates/財產分析書.docx"
    
    try:
        template = DocxTemplate(template_path)
        
        # 準備簡單的測試資料
        test_data = {
            "insurance_company": "測試保險公司",
            "insured_person": "測試被保險人",
            "policyholder": "測試要保人",
            "vehicle_type": "測試車輛類型",
            "gender_male": "☑",
            "gender_female": "□",
            "CHECK_1": "☑",
            "CHECK_2": "☑"
        }
        
        print("🔄 開始渲染模板...")
        template.render(test_data)
        
        # 儲存結果
        output_path = "property_reports/test_template_output.docx"
        os.makedirs("property_reports", exist_ok=True)
        template.save(output_path)
        
        print(f"✅ 模板渲染成功: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ 模板渲染失敗: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """主函數"""
    print("🧪 Word模板測試")
    print("=" * 60)
    
    # 測試基本功能
    basic_success = test_template_basic()
    docxtpl_success = test_docxtpl_basic()
    
    # 測試渲染
    render_success = False
    if basic_success and docxtpl_success:
        render_success = test_simple_render()
    
    print("\n📊 測試結果")
    print("=" * 60)
    print(f"基本讀取: {'✅ 成功' if basic_success else '❌ 失敗'}")
    print(f"docxtpl讀取: {'✅ 成功' if docxtpl_success else '❌ 失敗'}")
    print(f"模板渲染: {'✅ 成功' if render_success else '❌ 失敗'}")
    
    if basic_success and docxtpl_success and render_success:
        print("\n🎉 所有測試通過！Word模板功能正常")
    else:
        print("\n⚠️ 部分測試失敗，需要檢查模板檔案")

if __name__ == "__main__":
    main() 