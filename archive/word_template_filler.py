#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 模板填入系統
使用現有的 Word 檔案作為模板，填入辨識結果
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import os
import re
from datetime import datetime
from typing import Dict, List, Any, Optional
from docxtpl import DocxTemplate

class WordTemplateFiller:
    """Word 模板填入器"""
    
    def __init__(self, template_path: str):
        """
        初始化填入器
        
        Args:
            template_path: Word 模板檔案路徑
        """
        self.template_path = template_path
        self.document = None
        
        # 檢查模板檔案是否存在
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Word 模板檔案不存在: {template_path}")
    
    def load_template(self):
        """載入 Word 模板"""
        try:
            self.document = Document(self.template_path)
            print(f"✅ Word 模板載入成功: {self.template_path}")
            return True
        except Exception as e:
            print(f"❌ Word 模板載入失敗: {str(e)}")
            return False
    
    def fill_checkboxes(self, checkbox_map: Dict[str, bool]):
        """
        根據 checkbox_map 自動勾選 Word 內所有指定選項（用☑），支援分開 run
        Args:
            checkbox_map: { '保障需求': True, '否': True, ... }
        """
        if not self.document:
            self.load_template()
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        runs = para.runs
                        for i, run in enumerate(runs):
                            # 只處理 run 內容為「□」的情境
                            if run.text == '□' and i + 1 < len(runs):
                                next_text = runs[i+1].text.strip()
                                for key, checked in checkbox_map.items():
                                    if next_text.startswith(key):
                                        if checked:
                                            run.text = '☑'
                                        else:
                                            run.text = '□'

    def fill_checkboxes_single(self, single_choice_map: Dict[str, str]):
        """
        只勾同一行的單一選項（單選），根據指定的選項關鍵字
        Args:
            single_choice_map: { '保險類型': '財產保險', ... }
        """
        if not self.document:
            self.load_template()
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        runs = para.runs
                        # 將所有 run 內容合併成一行字串
                        full_text = ''.join([r.text for r in runs])
                        for topic, answer in single_choice_map.items():
                            # 只處理有該題目關鍵字的段落
                            if topic in full_text:
                                # 找到所有 run 為「□」的 index
                                for i, run in enumerate(runs):
                                    if run.text == '□' and i + 1 < len(runs):
                                        next_text = runs[i+1].text.strip()
                                        # 只勾指定答案
                                        if next_text.startswith(answer):
                                            run.text = '☑'
                                        else:
                                            run.text = '□'

    def fill_checkboxes_by_option(self, options_to_check: list):
        """
        只根據選項關鍵字打勾，遇到指定選項就打勾，其他不動
        Args:
            options_to_check: ['財產保險', '否', '是']
        """
        if not self.document:
            self.load_template()
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        runs = para.runs
                        for i, run in enumerate(runs):
                            if run.text == '□' and i + 1 < len(runs):
                                next_text = runs[i+1].text.strip()
                                for key in options_to_check:
                                    if next_text.startswith(key):
                                        run.text = '☑'

    def fill_insurance_data(self, insurance_data: Dict, output_path: str) -> str:
        """
        填入保險資料到 Word 模板（只處理打勾部分）
        
        Args:
            insurance_data: 保險資料字典
            output_path: 輸出檔案路徑
            
        Returns:
            生成的檔案路徑
        """
        if not self.document:
            if not self.load_template():
                return None
        
        try:
            print("🔄 開始填入資料到 Word 模板...")
            
            # 只根據選項關鍵字打勾
            options_to_check = [
                '財產保險',
                '否',  # 只要遇到「否」就打勾
                '是',  # 只要遇到「是」就打勾
            ]
            self.fill_checkboxes_by_option(options_to_check)
            
            # 儲存文件
            self.document.save(output_path)
            print(f"✅ Word 檔案生成成功: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ 填入資料失敗: {str(e)}")
            return None
    
    def _fill_insurance_type_section(self, data: Dict):
        """填入保險類型選擇區塊"""
        print("  📝 填入保險類型選擇...")
        
        # 根據資料判斷保險類型
        vehicle_type = data.get('vehicle_type', '')
        coverage_items = data.get('coverage_items', [])
        
        # 判斷是否為車險
        is_vehicle_insurance = any([
            '車' in vehicle_type,
            any('車' in str(item.get('保險種類', '')) for item in coverage_items),
            any('強制' in str(item.get('保險種類', '')) for item in coverage_items)
        ])
        
        # 在文件中尋找並填入
        self._replace_text_in_document("□財產保險", "✓財產保險" if is_vehicle_insurance else "□財產保險")
    
    def _fill_analysis_report_section(self, data: Dict):
        """填入財產保險契約分析報告書區塊"""
        print("  📝 填入財產保險契約分析報告書...")
        
        # 填入保障需求
        self._replace_text_in_document("□保障需求", "✓保障需求")
        
        # 填入是否指定保險公司（根據資料判斷）
        insurance_company = data.get('insurance_company', '')
        has_specified_company = bool(insurance_company and insurance_company.strip())
        self._replace_text_in_document("□否", "✓是" if has_specified_company else "✓否")
        
        # 填入是否已投保其他保險
        # 這裡可以根據實際需求調整邏輯
        self._replace_text_in_document("□否", "✓否")  # 預設為否
        
        # 填入投保目的
        self._replace_text_in_document("□保障需求", "✓保障需求")
        
        # 填入要保人確認事項
        self._replace_text_in_document("□是", "✓是")
    
    def _fill_broker_recommendations(self, data: Dict):
        """填入保險經紀人建議事項"""
        print("  📝 填入保險經紀人建議事項...")
        
        # 填入保障範圍
        vehicle_type = data.get('vehicle_type', '')
        if '車' in vehicle_type:
            self._replace_text_in_document("□車險", "✓車險")
        
        # 填入保險商品/名稱/保額
        coverage_items = data.get('coverage_items', [])
        if coverage_items:
            # 將承保內容轉換為字串格式
            coverage_text = []
            for item in coverage_items:
                if isinstance(item, dict):
                    coverage_text.append(item.get('保險種類', ''))
                else:
                    coverage_text.append(str(item))
            
            # 填入承保內容
            coverage_summary = '、'.join(filter(None, coverage_text))
            if coverage_summary:
                self._replace_text_in_document("□同上述之", "✓同上述之")
                # 可以進一步填入具體的承保內容
                self._replace_text_in_document("承保內容", coverage_summary)
        
        # 填入建議投保保險公司理由
        self._replace_text_in_document("□商品保障內容符合客戶需求", "✓商品保障內容符合客戶需求")
    
    def _fill_customer_info(self, data: Dict):
        """填入客戶基本資料"""
        print("  📝 填入客戶基本資料...")
        
        # 填入要保人資料
        policyholder = data.get('policyholder', '')
        if policyholder:
            self._replace_text_in_document("要保人姓名", policyholder)
        
        # 填入被保險人資料
        insured_person = data.get('insured_person', '')
        if insured_person:
            self._replace_text_in_document("被保險人姓名", insured_person)
        
        # 填入車牌號碼
        license_number = data.get('license_number', '')
        if license_number:
            self._replace_text_in_document("車牌號碼", license_number)
        
        # 填入保險公司
        insurance_company = data.get('insurance_company', '')
        if insurance_company:
            self._replace_text_in_document("保險公司", insurance_company)
    
    def _replace_text_in_document(self, old_text: str, new_text: str):
        """
        在文件中替換文字
        
        Args:
            old_text: 要替換的文字
            new_text: 新的文字
        """
        if not self.document:
            return
        
        # 在段落中替換
        for paragraph in self.document.paragraphs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
        
        # 在表格中替換
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(old_text, new_text)
    
    def _find_and_replace_in_tables(self, old_text: str, new_text: str):
        """
        在表格中尋找並替換文字
        
        Args:
            old_text: 要替換的文字
            new_text: 新的文字
        """
        if not self.document:
            return
        
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            # 保持原有格式，只替換文字
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
    
    def get_template_info(self):
        if not self.document:
            raise RuntimeError("尚未載入 Word 文件，請先呼叫 load_template()。")
        return {
            "paragraphs_count": len(self.document.paragraphs),
            "tables_count": len(self.document.tables),
            "sections_count": len(self.document.sections)
        }

    def _replace_tag_with_value(self, para, tag, value, debug=False):
        """
        跨 run 合併自動化：將連續 run.text 合起來等於 tag 的部分，合併為一個 run 並替換為 value，保留第一個 run 格式。
        debug: 若為 True，會印出 run 結構與比對過程
        """
        runs = para.runs
        n = len(runs)
        if debug:
            print("--- 段落內容 ---")
            print(f"原始段落: {para.text}")
            for idx, run in enumerate(runs):
                print(f"  run[{idx}]: '{run.text}'")
        i = 0
        while i < n:
            concat = ''
            for j in range(i, n):
                concat += runs[j].text
                if debug:
                    print(f"  嘗試合併 run[{i}~{j}]: '{concat}' (目標: '{tag}')")
                if concat == tag:
                    if debug:
                        print(f"  >> 命中！run[{i}~{j}] 合併為 '{value}'")
                    runs[i].text = str(value)
                    for k in range(i+1, j+1):
                        runs[k].text = ''
                    i = j
                    break
                elif not tag.startswith(concat):
                    break
            i += 1

    def fill_check_marks(self, check_tags=None, debug=False):
        """
        支援自動合併 run，將所有標記換成 ☑。
        Args:
            check_tags: 要打勾的標記 list，例如 ["{{CHECK_1}}", "{{CHECK_2}}"]
            debug: 是否開啟 debug 輸出
        """
        if check_tags is None:
            check_tags = ["{{CHECK_1}}", "{{CHECK_2}}"]
        if not self.document:
            raise RuntimeError("尚未載入 Word 文件，請先呼叫 load_template()。")
        # 處理段落
        for para in self.document.paragraphs:
            for tag in check_tags:
                self._replace_tag_with_value(para, tag, "☑", debug=debug)
        # 處理表格
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for tag in check_tags:
                            self._replace_tag_with_value(para, tag, "☑", debug=debug)

    def fill_fields(self, field_map: dict, debug=False):
        """
        支援自動合併 run，將所有標記換成對應值。
        Args:
            field_map: 欄位對應 dict，例如 {"{{insurance_company}}": "新安東京海上產物保險股份有限公司", ...}
            debug: 是否開啟 debug 輸出
        """
        if not self.document:
            raise RuntimeError("尚未載入 Word 文件，請先呼叫 load_template()。")
        # 處理段落
        for para in self.document.paragraphs:
            for tag, value in field_map.items():
                self._replace_tag_with_value(para, tag, value, debug=debug)
        # 處理表格
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for tag, value in field_map.items():
                            self._replace_tag_with_value(para, tag, value, debug=debug)

def main():
    """docxtpl 樣板自動填入完整欄位範例（含自動勾選性別）"""
    print("📄 Word 模板填入系統測試 (docxtpl 完整欄位+性別勾選)")
    print("=" * 50)
    template_path = "assets/templates/財產分析書.docx"
    output_path = "property_reports/test_docxtpl.docx"
    os.makedirs('property_reports', exist_ok=True)

    # 假設這是 OCR 辨識結果
    ocr_result = {
        "gender": "男",
        "policyholder_gender": "女",
        # ...其他欄位
    }

    context = {
        # === 自動勾選性別部分 ===
        "gender_male": "☑" if ocr_result.get("gender") == "男" else "□",
        "gender_female": "☑" if ocr_result.get("gender") == "女" else "□",
        "policyholder_gender_male": "☑" if ocr_result.get("policyholder_gender") == "男" else "□",
        "policyholder_gender_female": "☑" if ocr_result.get("policyholder_gender") == "女" else "□",
        # === 自動勾選關係選項 ===
        # 依據 ocr_result['relationship'] 自動決定哪個選項打勾
        **{
            f"CHECK_RELATIONSHIP_{opt}": ("☑" if ocr_result.get("relationship") == opt else "□")
            for opt in [
                "本人", "配偶", "父母", "子女", "雇傭", "祖孫", "債權債務", "標的物"
            ]
        },
        # === 固定勾選部分 ===
        "CHECK_1": "☑",
        "CHECK_2": "☑",
        # === OCR 欄位部分 ===
        "insurance_company": "新安東京海上產物保險股份有限公司",
        "insured_section": "被保險人區塊",
        "insured_person": "王小明",
        "legal_representative": "林經理",
        "id_number": "A123456789",
        "birth_date": "1990-01-01",
        "gender": ocr_result.get("gender", ""),
        "policyholder_section": "要保人區塊",
        "policyholder": "王小明",
        "relationship": "子女",
        "policyholder_legal_representative": "林經理",
        "policyholder_gender": ocr_result.get("policyholder_gender", ""),
        "policyholder_id": "A123456789",
        "policyholder_birth_date": "1990-01-01",
        "vehicle_type": "小客車",
        "license_number": "ABC-1234",
        "coverage_items": "車體險、強制險",
        "total_premium": "27,644",
        "compulsory_insurance_period": "",
        "optional_insurance_period": "自民國114年5月20日中午12時起至民國115年5月20日中午12時止",
        # ...如有其他欄位可再擴充
    }
    try:
        tpl = DocxTemplate(template_path)
        tpl.render(context)
        tpl.save(output_path)
        print(f"✅ docxtpl 完整欄位產生完成，請用 Word 開啟檢查: {output_path}")
    except Exception as e:
        print(f"❌ docxtpl 測試失敗: {str(e)}")

# === 第三塊：進階自動化（如表格、條件判斷等） ===
# 這裡可擴充 for 迴圈、if 判斷、表格自動產生等 Jinja2 樣板邏輯
# ...
# ... existing code ...

if __name__ == "__main__":
    main() 