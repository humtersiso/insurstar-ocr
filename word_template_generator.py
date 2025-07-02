#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 模板生成器
使用 python-docx 建立財產分析書 Word 檔案
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime
from typing import Dict, List, Any, Optional

class WordTemplateGenerator:
    """Word 模板生成器"""
    
    def __init__(self):
        """初始化生成器"""
        self.document = None
    
    def create_property_analysis_docx(self, insurance_data: Dict, output_path: str) -> str:
        """
        建立財產分析書 Word 檔案
        
        Args:
            insurance_data: 保險資料字典
            output_path: 輸出檔案路徑
            
        Returns:
            生成的檔案路徑
        """
        # 建立新文件
        self.document = Document()
        
        # 設定頁面邊距
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # 1. 標題
        self._add_title()
        
        # 2. 公司資訊
        self._add_company_info()
        
        # 3. 保險類型選擇
        self._add_insurance_type_section()
        
        # 4. 客戶基本資料
        self._add_customer_info_section(insurance_data)
        
        # 5. 財產保險契約分析報告書
        self._add_analysis_report_section(insurance_data)
        
        # 儲存文件
        self.document.save(output_path)
        return output_path
    
    def _add_title(self):
        """加入標題"""
        title = self.document.add_heading('創星保險經紀人股份有限公司分析報告書', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 設定標題樣式
        title_style = title.style
        title_style.font.size = Pt(16)
        title_style.font.bold = True
    
    def _add_company_info(self):
        """加入公司資訊"""
        # 加入空行
        self.document.add_paragraph()
        
        # 公司資訊表格
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # 設定欄寬
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(4)
        
        # 填入資料
        row = table.rows[0]
        row.cells[0].text = "創星保險經紀人(股)公司"
        row.cells[1].text = "總公司地址：台北市中山區民權東路二段46號3樓之1"
        
        # 設定樣式
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].style.font.size = Pt(10)
    
    def _add_insurance_type_section(self):
        """加入保險類型選擇區塊"""
        # 加入空行
        self.document.add_paragraph()
        
        # 保險類型表格
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # 設定欄寬
        for i in range(3):
            table.columns[i].width = Inches(2.5)
        
        # 填入資料
        row = table.rows[0]
        row.cells[0].text = "□人身保險"
        row.cells[1].text = "□財產保險"
        row.cells[2].text = "□旅行平安保險"
        
        # 設定樣式
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].style.font.size = Pt(10)
    
    def _add_customer_info_section(self, data: Dict):
        """加入客戶基本資料區塊"""
        # 加入空行
        self.document.add_paragraph()
        
        # 要保人資料
        policyholder_name = data.get('policyholder', '')
        policyholder_rep = data.get('policyholder_legal_representative', '')
        policyholder_gender = data.get('policyholder_gender', '')
        policyholder_id = data.get('policyholder_id', '')
        policyholder_birth = data.get('policyholder_birth_date', '')
        
        # 被保險人資料
        insured_name = data.get('insured_person', '')
        insured_rep = data.get('legal_representative', '')
        insured_gender = data.get('gender', '')
        insured_id = data.get('id_number', '')
        insured_birth = data.get('birth_date', '')
        
        # 車輛資料
        license_number = data.get('license_number', '')
        vehicle_type = data.get('vehicle_type', '')
        relationship = data.get('relationship', '')
        
        # 建立客戶資料表格
        table = self.document.add_table(rows=15, cols=5)
        table.style = 'Table Grid'
        
        # 設定欄寬
        col_widths = [1.5, 2.5, 2.5, 1.5, 2.0]  # 英吋
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width)
        
        # 填入資料
        rows = table.rows
        
        # 標題行
        rows[0].cells[0].text = "一、客戶基本資料"
        rows[0].cells[0].merge(rows[0].cells[4])
        rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rows[0].cells[0].paragraphs[0].style.font.bold = True
        
        # 要保人資料
        rows[2].cells[0].text = "要保人"
        rows[2].cells[1].text = "姓名/公司行號"
        rows[2].cells[2].text = policyholder_name
        rows[2].cells[3].text = "法人代表人"
        rows[2].cells[4].text = policyholder_rep
        
        # 性別選擇
        gender_text = "✓" if policyholder_gender == "男" else "□" if policyholder_gender == "女" else "□"
        rows[3].cells[1].text = "□男 □女"
        rows[3].cells[2].text = gender_text
        
        # 身分證和出生日期
        rows[4].cells[1].text = "身分證字號/統編"
        rows[4].cells[2].text = policyholder_id
        rows[4].cells[3].text = "出生年月日"
        rows[4].cells[4].text = policyholder_birth
        
        # 職業
        rows[5].cells[1].text = "職業"
        
        # 被保險人資料
        rows[7].cells[0].text = "被保險人"
        rows[7].cells[1].text = "姓名/公司行號"
        rows[7].cells[2].text = insured_name
        rows[7].cells[3].text = "法人代表人"
        rows[7].cells[4].text = insured_rep
        
        # 性別選擇
        gender_text = "✓" if insured_gender == "男" else "□" if insured_gender == "女" else "□"
        rows[8].cells[1].text = "□男 □女"
        rows[8].cells[2].text = gender_text
        
        # 身分證和出生日期
        rows[9].cells[1].text = "身分證字號/統編"
        rows[9].cells[2].text = insured_id
        rows[9].cells[3].text = "出生年月日"
        rows[9].cells[4].text = insured_birth
        
        # 職業
        rows[10].cells[1].text = "職業"
        
        # 車輛資料
        rows[12].cells[0].text = "投保車險必填"
        rows[12].cells[1].text = "車牌號碼"
        rows[12].cells[2].text = license_number
        
        # 車輛類型
        vehicle_check = "✓" if "汽車" in vehicle_type else "✓" if "機車" in vehicle_type else "□"
        rows[13].cells[1].text = "□汽車 □機車"
        rows[13].cells[2].text = vehicle_check
        
        # 關係
        rows[14].cells[1].text = "要／被保險人關係"
        rows[14].cells[2].text = relationship
        
        # 設定表格樣式
        for row in rows:
            for cell in row.cells:
                cell.paragraphs[0].style.font.size = Pt(9)
    
    def _add_analysis_report_section(self, data: Dict):
        """加入財產保險契約分析報告書區塊"""
        # 加入空行
        self.document.add_paragraph()
        
        # 取得承保內容
        coverage_items = data.get('coverage_items', [])
        total_premium = data.get('total_premium', '')
        
        # 建立分析報告表格
        table_data = [
            ["二、財產保險契約分析報告書", "", "", ""],
            ["", "", "", ""],
            ["本次投保之目的", "", "", ""],
            ["", "", "", ""],
            ["欲投保之保險種類", "保險金額", "自負額", "簽單保費"],
        ]
        
        # 加入承保內容
        for item in coverage_items:
            if isinstance(item, dict):
                insurance_type = item.get('保險種類', '')
                insurance_amount = item.get('保險金額', '')
                deductible = item.get('自負額', '')
                premium = item.get('簽單保費', '')
                
                table_data.append([insurance_type, insurance_amount, deductible, premium])
                
                # 加入子項目
                sub_items = item.get('sub_items', [])
                for sub_item in sub_items:
                    if isinstance(sub_item, dict):
                        sub_type = sub_item.get('保險種類', '')
                        sub_amount = sub_item.get('保險金額', '')
                        sub_deductible = sub_item.get('自負額', '')
                        sub_premium = sub_item.get('簽單保費', '')
                        
                        table_data.append([f"  └ {sub_type}", sub_amount, sub_deductible, sub_premium])
        
        # 加入總計
        table_data.append(["", "", "總保險費", total_premium])
        
        # 建立表格
        table = self.document.add_table(rows=len(table_data), cols=4)
        table.style = 'Table Grid'
        
        # 設定欄寬
        col_widths = [4.0, 2.0, 2.0, 2.0]  # 英吋
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width)
        
        # 填入資料
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = str(cell_data)
                row.cells[j].paragraphs[0].style.font.size = Pt(9)
        
        # 設定標題行樣式
        title_row = table.rows[0]
        title_row.cells[0].merge(title_row.cells[3])
        title_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_row.cells[0].paragraphs[0].style.font.bold = True
        
        # 設定欄位標題行樣式
        header_row = table.rows[4]
        for cell in header_row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].style.font.bold = True

def main():
    """測試函數"""
    print("📄 Word 模板生成器測試")
    print("=" * 50)
    
    # 初始化生成器
    generator = WordTemplateGenerator()
    
    # 測試資料
    test_data = {
        'policyholder': '張三',
        'policyholder_legal_representative': '',
        'policyholder_gender': '男',
        'policyholder_id': 'A123456789',
        'policyholder_birth_date': '1980/01/01',
        'insured_person': '張三',
        'legal_representative': '',
        'gender': '男',
        'id_number': 'A123456789',
        'birth_date': '1980/01/01',
        'license_number': 'ABC-1234',
        'vehicle_type': '汽車',
        'relationship': '本人',
        'coverage_items': [
            {
                '保險代號': '05',
                '保險種類': '車體損失保險乙式(Q)',
                '保險金額': '',
                '自負額': '20,000',
                '簽單保費': '20,527',
                'sub_items': [
                    {'保險種類': '每一個人傷害', '保險金額': '300萬', '自負額': '', '簽單保費': '9,485'},
                    {'保險種類': '每一意外事故之傷害', '保險金額': '600萬', '自負額': '', '簽單保費': ''}
                ]
            }
        ],
        'total_premium': '29,012'
    }
    
    # 生成 Word 檔案
    output_path = 'outputs/test_property_analysis.docx'
    os.makedirs('outputs', exist_ok=True)
    
    try:
        result_path = generator.create_property_analysis_docx(test_data, output_path)
        print(f"✅ Word 檔案生成成功: {result_path}")
        print("📝 請用 Microsoft Word 或 LibreOffice 開啟檢查內容")
    except Exception as e:
        print(f"❌ 生成失敗: {str(e)}")

if __name__ == "__main__":
    main() 