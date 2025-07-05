#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 模板處理系統
讀取創星保經財產分析書模板，填入OCR辨識結果，處理勾選標記
"""

import os
import re
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docxtpl import DocxTemplate, InlineImage
import json

class WordTemplateProcessorPure:
    """純 docxtpl context 版 Word 模板處理器"""
    def __init__(self, template_path: Optional[str] = None):
        if template_path is None:
            template_path = "assets/templates/財產分析書.docx"
        self.template_path = template_path
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Word 模板檔案不存在: {template_path}")
        # OCR欄位對照表
        self.field_mapping = {
            "insurance_company": "{{insurance_company}}",
            "insured_section": "{{insured_section}}",
            "insured_person": "{{insured_person}}",
            "legal_representative": "{{legal_representative}}",
            "id_number": "{{id_number}}",
            "birth_date": "{{birth_date}}",
            "gender": "{{gender}}",
            "policyholder_section": "{{policyholder_section}}",
            "policyholder": "{{policyholder}}",
            "relationship": "{{relationship}}",
            "policyholder_legal_representative": "{{policyholder_legal_representative}}",
            "policyholder_gender": "{{policyholder_gender}}",
            "policyholder_id": "{{policyholder_id}}",
            "policyholder_birth_date": "{{policyholder_birth_date}}",
            "vehicle_type": "{{vehicle_type}}",
            "license_number": "{{license_number}}",
            "coverage_items": "{{coverage_items}}",
            "total_premium": "{{total_premium}}",
            "compulsory_insurance_period": "{{compulsory_insurance_period}}",
            "optional_insurance_period": "{{optional_insurance_period}}",
            "optional_insurance_amount": "{{optional_insurance_amount}}"
        }
        self.checkbox_mapping = {
            "gender_male": "{{gender_male}}",
            "gender_female": "{{gender_female}}",
            "policyholder_gender_male": "{{policyholder_gender_male}}",
            "policyholder_gender_female": "{{policyholder_gender_female}}",
            "relationship_self": "{{CHECK_RELATIONSHIP_本人}}",
            "relationship_spouse": "{{CHECK_RELATIONSHIP_配偶}}",
            "relationship_parent": "{{CHECK_RELATIONSHIP_父母}}",
            "relationship_child": "{{CHECK_RELATIONSHIP_子女}}",
            "relationship_employee": "{{CHECK_RELATIONSHIP_雇傭}}",
            "relationship_grandparent": "{{CHECK_RELATIONSHIP_祖孫}}",
            "relationship_creditor": "{{CHECK_RELATIONSHIP_債權債務}}",
            "relationship_object": "{{CHECK_RELATIONSHIP_標的物}}",
            "check_1": "{{CHECK_1}}",
            "check_2": "{{CHECK_2}}"
        }

    def set_checkbox_font(self, docx_path: str, extra_fields: dict = {}):
        """設定勾選框為粗體"""
        doc = Document(docx_path)
        
        # 只設定勾選框為粗體
        for para in doc.paragraphs:
            for run in para.runs:
                if "☑" in run.text:
                    run.bold = True
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if "☑" in run.text:
                                run.bold = True
        
        doc.save(docx_path)

    def process_ocr_data(self, ocr_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        處理OCR辨識資料，轉換為Word模板所需的格式
        
        Args:
            ocr_data: OCR辨識結果字典
            
        Returns:
            處理後的資料字典
        """
        # 將所有值為 '無填寫' 的欄位轉為空字串
        ocr_data = {k: ("" if isinstance(v, str) and v.strip() == "無填寫" else v) for k, v in ocr_data.items()}

        processed_data = {}
        
        # 處理基本欄位
        for field, template_tag in self.field_mapping.items():
            if field in ["compulsory_insurance_period", "optional_insurance_period"]:
                # 若為空值則保留原本內容
                value = ocr_data.get(field)
                if value is None:
                    processed_data[field] = None
                else:
                    processed_data[field] = str(value)
            elif field == "total_premium":
                # total_premium 直接使用原始值，不進行額外處理
                value = ocr_data.get(field, "")
                processed_data[field] = str(value) if value else ""
            else:
                value = ocr_data.get(field, "")
                processed_data[field] = str(value) if value else ""
        
        # 性別勾選（被保險人）
        gender = ocr_data.get("gender", "")
        if gender == "男":
            processed_data["gender_male"] = "☑ "
            processed_data["gender_female"] = "□"
        elif gender == "女":
            processed_data["gender_male"] = "□"
            processed_data["gender_female"] = "☑ "
        else:
            processed_data["gender_male"] = "□"
            processed_data["gender_female"] = "□"
        processed_data["gender"] = gender
        
        # 性別勾選（要保人）
        policyholder_gender = ocr_data.get("policyholder_gender", "")
        if policyholder_gender == "男":
            processed_data["policyholder_gender_male"] = "☑ "
            processed_data["policyholder_gender_female"] = "□"
        elif policyholder_gender == "女":
            processed_data["policyholder_gender_male"] = "□"
            processed_data["policyholder_gender_female"] = "☑ "
        else:
            processed_data["policyholder_gender_male"] = "□"
            processed_data["policyholder_gender_female"] = "□"
        processed_data["policyholder_gender"] = policyholder_gender
        
        # 車種勾選
        vehicle_type = ocr_data.get("vehicle_type", "")
        if "機車" in vehicle_type:
            processed_data["vehicle_type_moto"] = "☑ "
            processed_data["vehicle_type_car"] = "□"
        else:
            processed_data["vehicle_type_moto"] = "□"
            processed_data["vehicle_type_car"] = "☑ "
        processed_data["vehicle_type"] = vehicle_type
        
        # 關係勾選（使用正確的標記名稱）
        relationship = ocr_data.get("relationship", "")
        
        # 新版關係勾選（relationship_1 ~ relationship_8）
        relationship_map = {
            "本人": "relationship_1",
            "配偶": "relationship_2",
            "父母": "relationship_3",
            "子女": "relationship_4",
            "雇傭": "relationship_5",
            "祖孫": "relationship_6",
            "債權債務": "relationship_7",
            "標的物": "relationship_8"
        }
        for rel, tag in relationship_map.items():
            processed_data[tag] = "☑ " if relationship == rel else "□"
        
        # 舊版關係勾選（保留原有 CHECK_RELATIONSHIP_XXX 標記）
        relationship_options = [
            "本人", "配偶", "父母", "子女", "雇傭", "祖孫", "債權債務", "標的物"
        ]
        for option in relationship_options:
            key = f"CHECK_RELATIONSHIP_{option}"
            processed_data[key] = "☑ " if relationship == option else "□"
        
        processed_data["relationship"] = relationship
        
        # 固定勾選項目
        processed_data["CHECK_1"] = "☑ "
        processed_data["CHECK_2"] = "☑ "
        
        # 保期勾選
        compulsory_period = ocr_data.get("compulsory_insurance_period", None)
        optional_period = ocr_data.get("optional_insurance_period", None)
        processed_data["check_compulsory_insurance_period"] = "☑ " if compulsory_period not in [None, ""] else "□"
        processed_data["check_optional_insurance_period"] = "☑ " if optional_period not in [None, ""] else "□"
        
        # 空格補齊：若欄位為空，補空格
        for field in ["policyholder_gender", "relationship", "gender", "vehicle_type"]:
            if not processed_data.get(field):
                processed_data[field] = " "
        
        # 判斷強制/任意險有無
        compulsory = ocr_data.get("compulsory_insurance_period")
        optional = ocr_data.get("optional_insurance_period")
        coverage_items = ocr_data.get("coverage_items", [])
        def find_car_damage_amount():
            for item in coverage_items:
                if "車體損失保險" in item.get("保險種類", ""):
                    return item.get("保險金額", "")
            return ""
        def find_third_party_personal_amount():
            for item in coverage_items:
                if "第三人傷害責任險" in item.get("保險種類", ""):
                    for sub in item.get("sub_items", []):
                        if "每一個人傷害" in sub.get("保險種類", ""):
                            return sub.get("保險金額", "")
            return ""
        optional_insurance_amount = ""
        if compulsory and not optional:
            # 只有強制險，金額不填
            pass
        elif optional:
            # 有任意險（不論強制險有無）
            car_damage = find_car_damage_amount()
            if car_damage:
                optional_insurance_amount = car_damage
            else:
                optional_insurance_amount = find_third_party_personal_amount()
        
        # 移除「萬」字
        if optional_insurance_amount:
            optional_insurance_amount = optional_insurance_amount.replace("萬", "").strip()
        
        processed_data["optional_insurance_amount"] = optional_insurance_amount
        
        # 若有需要填入強制險金額欄位，這裡不填值（保留空白）
        processed_data["compulsory_insurance_amount"] = ""
        
        return processed_data

    def fill_template(self, ocr_data: Dict[str, Any], output_path: Optional[str] = None) -> Optional[str]:
        """
        填入OCR資料到Word模板，並自動補齊所有欄位
        """
        try:
            print("🔄 開始處理Word模板...")
            processed_data = self.process_ocr_data(ocr_data)
            tpl = DocxTemplate(self.template_path)
            # 圖片欄位直接用 InlineImage
            processed_data['watermark_name_blue'] = InlineImage(tpl, 'assets/watermark/watermark_name_blue.png', width=Mm(37))
            processed_data['watermark_company_blue'] = InlineImage(tpl, 'assets/watermark/watermark_company_blue.png', width=Mm(37))
            # PCN 欄位寫死
            processed_data['PCN'] = 'BB2H699299'
            if output_path is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"property_reports/財產分析書_{timestamp}.docx"
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            tpl.render(processed_data)
            tpl.save(output_path)
            print(f"✅ Word 檔案生成成功: {output_path}")
            self.set_checkbox_font(output_path)
            return output_path
        except Exception as e:
            print(f"❌ 填入資料失敗: {str(e)}")
            return None 